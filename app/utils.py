"""Shared utilities for document conversions."""
from __future__ import annotations

import io
import tempfile
from pathlib import Path

from pdf2docx import Converter
from docx import Document
from docx.shared import Inches


def pdf_to_docx_bytes(pdf_bytes: bytes, pdf_filename: str) -> tuple[bytes, str]:
    """Convert a single PDF (as bytes) to DOCX bytes and return output name."""
    stem = Path(pdf_filename).stem or "document"
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_tmp:
        pdf_tmp.write(pdf_bytes)
        pdf_path = Path(pdf_tmp.name)
    docx_path = pdf_path.with_suffix(".docx")

    try:
        converter = Converter(str(pdf_path))
        converter.convert(str(docx_path))
        converter.close()
        docx_bytes = docx_path.read_bytes()
    finally:
        pdf_path.unlink(missing_ok=True)
        docx_path.unlink(missing_ok=True)

    return docx_bytes, f"{stem}.docx"


def images_to_docx_bytes(images: list, docx_name: str = "images_to_word.docx") -> tuple[bytes, str]:
    """Embed multiple images into a DOCX and return bytes with file name."""
    document = Document()
    for uploaded in images:
        document.add_heading(uploaded.name, level=2)
        image_bytes = uploaded.read()
        stream = io.BytesIO(image_bytes)
        document.add_picture(stream, width=Inches(6))
        document.add_page_break()

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read(), docx_name
